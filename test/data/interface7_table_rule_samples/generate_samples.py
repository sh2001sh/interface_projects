from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


BASE_DIR = Path(__file__).resolve().parent

HEADERS = [
    "源协议",
    "源消息",
    "目标协议",
    "目标消息",
    "目标字段",
    "源字段",
    "转换关系",
    "转换类型",
    "字段含义",
    "说明",
]

TEMP_ROWS = [
    [
        "Temp_Sensor",
        "Temp_Sensor",
        "Temp_Report",
        "Temp_Report",
        "temperature_c",
        "temperature",
        "temperature_c = temperature；自然语言：源温度已经是摄氏度，目标温度字段直接继承源字段。",
        "transcoding",
        "温度",
        "用于验证接口7能从表格中抽取直接映射规则，并可继续交给接口8生成工程。",
    ],
    [
        "Temp_Sensor",
        "Temp_Sensor",
        "Temp_Report",
        "Temp_Report",
        "alarm",
        "status",
        "alarm = 0->0, 1->1, 2->1, 3->2；自然语言：状态码0表示正常，1和2都转为一般告警，3转为严重告警。",
        "mapping",
        "告警等级",
        "用于验证枚举/状态映射关系。",
    ],
]

K16_ROWS = [
    [
        "K1_6",
        "K1.6",
        "X0_5",
        "X0.5",
        "时间1",
        "小时",
        "时间1 = 小时；自然语言：目标时间1使用源消息中的小时字段。",
        "transcoding",
        "时间小时",
        "拆分时间字段，单源字段直传。",
    ],
    [
        "K1_6",
        "K1.6",
        "X0_5",
        "X0.5",
        "时间2",
        "分钟",
        "时间2 = 分钟；自然语言：目标时间2使用源消息中的分钟字段。",
        "transcoding",
        "时间分钟",
        "拆分时间字段，单源字段直传。",
    ],
    [
        "K1_6",
        "K1.6",
        "X0_5",
        "X0.5",
        "时间3",
        "秒",
        "时间3 = 秒；自然语言：目标时间3使用源消息中的秒字段。",
        "transcoding",
        "时间秒",
        "拆分时间字段，单源字段直传。",
    ],
    [
        "K1_6",
        "K1.6",
        "X0_5",
        "X0.5",
        "时间4",
        "日,小时,分钟,秒",
        "时间4 = (日 * 86400) + (小时 * 3600) + (分钟 * 60) + 秒；自然语言：将日、小时、分钟、秒合成为秒级时间累计值。",
        "transcoding",
        "累计秒",
        "多源字段算术转换。",
    ],
    [
        "K1_6",
        "K1.6",
        "X0_5",
        "X0.5",
        "高度",
        "高程",
        "高度 = 高程；自然语言：源协议高程与目标高度语义一致，直接写入。",
        "transcoding",
        "高度",
        "字段别名映射。",
    ],
    [
        "K1_6",
        "K1.6",
        "X0_5",
        "X0.5",
        "经度",
        "经度",
        "经度 = 经度；自然语言：目标经度直接沿用源经度。",
        "transcoding",
        "经度",
        "坐标字段直传。",
    ],
    [
        "K1_6",
        "K1.6",
        "X0_5",
        "X0.5",
        "纬度",
        "纬度",
        "纬度 = 纬度；自然语言：目标纬度直接沿用源纬度。",
        "transcoding",
        "纬度",
        "坐标字段直传。",
    ],
    [
        "K1_6",
        "K1.6",
        "X0_5",
        "X0.5",
        "运行状态",
        "威胁类型",
        "运行状态 = 1 if 威胁类型 > 0 else 0；自然语言：只要源威胁类型有效，目标网络运行状态置为运行。",
        "transcoding",
        "运行状态",
        "条件表达式规则。",
    ],
    [
        "K1_6",
        "K1.6",
        "X0_5",
        "X0.5",
        "横滚角",
        "翻滚1",
        "横滚角 = 翻滚1；自然语言：源翻滚角字段写入目标横滚角。",
        "transcoding",
        "姿态横滚",
        "姿态字段别名映射。",
    ],
    [
        "K1_6",
        "K1.6",
        "X0_5",
        "X0.5",
        "俯仰角",
        "俯仰1",
        "俯仰角 = 俯仰1；自然语言：源俯仰角字段写入目标俯仰角。",
        "transcoding",
        "姿态俯仰",
        "姿态字段别名映射。",
    ],
    [
        "K1_6",
        "K1.6",
        "X0_5",
        "X0.5",
        "航向",
        "偏航1",
        "航向 = 偏航1；自然语言：源偏航角字段写入目标航向字段。",
        "transcoding",
        "姿态航向",
        "姿态字段别名映射。",
    ],
]


def apply_sheet_style(sheet) -> None:
    header_fill = PatternFill("solid", fgColor="D9EAF7")
    for cell in sheet[1]:
        cell.font = Font(bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(vertical="center", wrap_text=True)
    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    widths = [14, 12, 14, 12, 14, 22, 58, 14, 16, 42]
    for index, width in enumerate(widths, start=1):
        sheet.column_dimensions[get_column_letter(index)].width = width
    sheet.freeze_panes = "A2"


def add_sheet(workbook: Workbook, title: str, rows: list[list[str]]) -> None:
    sheet = workbook.create_sheet(title)
    sheet.append(HEADERS)
    for row in rows:
        sheet.append(row)
    apply_sheet_style(sheet)


def write_excel(path: Path) -> None:
    workbook = Workbook()
    default = workbook.active
    workbook.remove(default)
    add_sheet(workbook, "temp_sensor_to_report", TEMP_ROWS)
    add_sheet(workbook, "k16_to_x05", K16_ROWS)
    workbook.save(path)


def add_docx_table(document: Document, title: str, rows: list[list[str]]) -> None:
    document.add_heading(title, level=2)
    table = document.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"
    for index, header in enumerate(HEADERS):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def write_docx(path: Path) -> None:
    document = Document()
    document.add_heading("接口7表格模式自然语言转换关系样例", level=1)
    document.add_paragraph(
        "这些表格用于 input_mode=table_rule。转换关系列保留自然语言说明，同时给出明确等式或映射，便于接口7抽取为规则。"
    )
    add_docx_table(document, "Temp_Sensor 到 Temp_Report", TEMP_ROWS)
    add_docx_table(document, "K1.6 到 X0.5", K16_ROWS)
    document.save(path)


def write_csv(path: Path, rows: list[list[str]]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(HEADERS)
        writer.writerows(rows)


def write_request_payloads() -> None:
    payloads = {
        "request_table_rule_docx.json": {
            "input_mode": "table_rule",
            "project_name": "interface7_table_rule_docx_sample",
            "table_rule_files": [str(BASE_DIR / "interface7_table_rule_samples.docx")],
            "rules_output_dir": str(BASE_DIR / "output"),
        },
        "request_table_rule_xlsx.json": {
            "input_mode": "table_rule",
            "project_name": "interface7_table_rule_xlsx_sample",
            "table_rule_files": [str(BASE_DIR / "interface7_table_rule_samples.xlsx")],
            "rules_output_dir": str(BASE_DIR / "output"),
        },
        "request_temp_sensor_csv.json": {
            "input_mode": "table_rule",
            "project_name": "temp_sensor_to_temp_report_table_rule",
            "table_rule_files": [str(BASE_DIR / "temp_sensor_to_report_rules.csv")],
            "source_protocol_type": "Temp_Sensor",
            "source_message_code": "Temp_Sensor",
            "target_protocol_type": "Temp_Report",
            "target_message_code": "Temp_Report",
            "rules_output_dir": str(BASE_DIR / "output"),
        },
    }
    for filename, payload in payloads.items():
        (BASE_DIR / filename).write_text(
            json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )


def main() -> None:
    write_excel(BASE_DIR / "interface7_table_rule_samples.xlsx")
    write_docx(BASE_DIR / "interface7_table_rule_samples.docx")
    write_csv(BASE_DIR / "temp_sensor_to_report_rules.csv", TEMP_ROWS)
    write_csv(BASE_DIR / "k16_to_x05_rules.csv", K16_ROWS)
    write_request_payloads()


if __name__ == "__main__":
    main()
