from __future__ import annotations

from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Tuple

from protocol_extractor import enrich_protocol_metadata


DEFAULT_QUICK_SAMPLE_UNITS = 6
MAX_QUICK_SAMPLE_UNITS = 16
DEFAULT_TEXT_CHAR_LIMIT = 12000
DOCX_TABLE_SAMPLE_ROWS = 20
EXCEL_SAMPLE_ROWS = 30
TABLE_CELL_LIMIT = 12


def _emit_progress(
    progress_callback: Optional[Callable[[dict], None]],
    *,
    stage: str,
    message: str,
    progress: float,
    **extra: Any,
) -> None:
    if progress_callback is None:
        return
    payload = {
        "stage": stage,
        "message": message,
        "progress": round(float(progress), 4),
    }
    payload.update({key: value for key, value in extra.items() if value is not None})
    progress_callback(payload)


def _normalize_sample_limit(value: Any, default: int = DEFAULT_QUICK_SAMPLE_UNITS) -> int:
    try:
        normalized = int(str(value).strip()) if value not in (None, "") else int(default)
    except (TypeError, ValueError):
        normalized = int(default)
    return max(1, min(normalized, MAX_QUICK_SAMPLE_UNITS))


def build_sample_indexes(total_units: int, sample_limit: int) -> List[int]:
    if total_units <= 0:
        return []
    resolved_limit = min(total_units, _normalize_sample_limit(sample_limit))
    if resolved_limit >= total_units:
        return list(range(total_units))
    if resolved_limit == 1:
        return [0]
    indexes = {
        int(round(idx * (total_units - 1) / (resolved_limit - 1)))
        for idx in range(resolved_limit)
    }
    return sorted(indexes)


def _truncate_text(text: str, char_limit: int = DEFAULT_TEXT_CHAR_LIMIT) -> str:
    normalized = str(text or "").strip()
    if len(normalized) <= char_limit:
        return normalized
    return normalized[:char_limit].rstrip()


def _count_protocol_fields(blocks: List[Dict[str, Any]]) -> int:
    return sum(len(block.get("metadata", {}).get("protocol_fields", [])) for block in blocks)


def _build_metrics(
    *,
    total_units: int,
    blocks: List[Dict[str, Any]],
    sampled_units: int,
    sampled_unit_indexes: List[int],
    analysis_mode: str = "quick",
) -> Dict[str, Any]:
    non_empty_blocks = [block for block in blocks if str(block.get("content", "")).strip()]
    readable_chars = sum(len(str(block.get("content", "")).strip()) for block in non_empty_blocks)
    return {
        "total_units": total_units,
        "total_blocks": len(blocks),
        "readable_blocks": len(non_empty_blocks),
        "readable_chars": readable_chars,
        "protocol_field_count": _count_protocol_fields(blocks),
        "analysis_mode": analysis_mode,
        "sampled_units": sampled_units,
        "sampled_unit_indexes": sampled_unit_indexes,
    }


def _sample_pdf_content(
    file_path: str,
    *,
    progress_callback: Optional[Callable[[dict], None]] = None,
    sample_unit_limit: int = DEFAULT_QUICK_SAMPLE_UNITS,
) -> Tuple[Dict[str, Any], List[str]]:
    import pdfplumber

    issues: List[str] = []
    blocks: List[Dict[str, Any]] = []
    with pdfplumber.open(file_path) as pdf:
        total_pages = len(pdf.pages)
        sample_indexes = build_sample_indexes(total_pages, sample_unit_limit)
        for offset, page_index in enumerate(sample_indexes, start=1):
            page = pdf.pages[page_index]
            text = _truncate_text(page.extract_text() or "")
            if text:
                block = {
                    "page_num": page_index + 1,
                    "content": text,
                    "type": "text",
                    "metadata": {},
                }
                blocks.append(enrich_protocol_metadata(block, False))
            _emit_progress(
                progress_callback,
                stage="quick_sampling_pdf",
                message=f"已采样 PDF 页 {offset}/{len(sample_indexes)}",
                progress=15.0 + (offset / max(len(sample_indexes), 1)) * 65.0,
                processed_pages=offset,
                total_pages=len(sample_indexes),
            )

    if total_pages <= 0:
        issues.append("未识别到有效页数")
    if not blocks:
        issues.append("采样页未提取到可读文本")
    return _build_metrics(
        total_units=total_pages,
        blocks=blocks,
        sampled_units=len(sample_indexes),
        sampled_unit_indexes=[index + 1 for index in sample_indexes],
    ), issues


def _iter_docx_sample_blocks(file_path: str, sample_unit_limit: int) -> Tuple[int, List[Dict[str, Any]]]:
    from docx import Document

    doc = Document(file_path)
    sample_blocks: List[Dict[str, Any]] = []
    total_units = 0

    for element in doc.element.body:
        if element.tag.endswith("tbl"):
            total_units += 1
            if len(sample_blocks) >= sample_unit_limit:
                continue
            for table in doc.tables:
                if table._tbl != element:
                    continue
                rows: List[str] = []
                for row in table.rows[:DOCX_TABLE_SAMPLE_ROWS]:
                    cells = [str(cell.text or "").strip() for cell in row.cells[:TABLE_CELL_LIMIT]]
                    if any(cells):
                        rows.append(" | ".join(cell for cell in cells if cell))
                content = _truncate_text("\n".join(rows))
                if content:
                    sample_blocks.append(
                        enrich_protocol_metadata(
                            {
                                "page_num": total_units,
                                "content": content,
                                "type": "table",
                                "metadata": {},
                            },
                            False,
                        )
                    )
                break
        elif element.tag.endswith("p"):
            for para in doc.paragraphs:
                if para._p != element:
                    continue
                text = _truncate_text(para.text)
                if text:
                    total_units += 1
                    if len(sample_blocks) < sample_unit_limit:
                        sample_blocks.append(
                            enrich_protocol_metadata(
                                {
                                    "page_num": total_units,
                                    "content": text,
                                    "type": "text",
                                    "metadata": {},
                                },
                                False,
                            )
                        )
                break

    return total_units, sample_blocks


def _sample_docx_content(
    file_path: str,
    *,
    progress_callback: Optional[Callable[[dict], None]] = None,
    sample_unit_limit: int = DEFAULT_QUICK_SAMPLE_UNITS,
) -> Tuple[Dict[str, Any], List[str]]:
    issues: List[str] = []
    resolved_limit = _normalize_sample_limit(sample_unit_limit)
    total_units, blocks = _iter_docx_sample_blocks(file_path, resolved_limit)
    _emit_progress(
        progress_callback,
        stage="quick_sampling_docx",
        message="已完成 DOCX 采样校验",
        progress=80.0,
        processed_blocks=len(blocks),
        total_blocks=len(blocks),
    )
    if total_units <= 0:
        issues.append("未识别到有效段落或表格")
    if not blocks:
        issues.append("采样段落未提取到可读文本")
    return _build_metrics(
        total_units=total_units,
        blocks=blocks,
        sampled_units=len(blocks),
        sampled_unit_indexes=[int(block.get("page_num", 0) or 0) for block in blocks],
    ), issues


def _sample_excel_content(
    file_path: str,
    ext: str,
    *,
    progress_callback: Optional[Callable[[dict], None]] = None,
    sample_unit_limit: int = DEFAULT_QUICK_SAMPLE_UNITS,
) -> Tuple[Dict[str, Any], List[str]]:
    issues: List[str] = []
    blocks: List[Dict[str, Any]] = []

    if ext == ".xlsx":
        from openpyxl import load_workbook

        workbook = load_workbook(file_path, data_only=True, read_only=True)
        try:
            sheet_names = list(workbook.sheetnames)
            sample_indexes = build_sample_indexes(len(sheet_names), sample_unit_limit)
            for offset, sheet_index in enumerate(sample_indexes, start=1):
                sheet = workbook[sheet_names[sheet_index]]
                rows: List[str] = []
                for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                    if row_index > EXCEL_SAMPLE_ROWS:
                        break
                    cleaned = [str(cell).strip() for cell in row[:TABLE_CELL_LIMIT] if str(cell or "").strip()]
                    if cleaned:
                        rows.append(" | ".join(cleaned))
                content = _truncate_text("\n".join(rows))
                if content:
                    blocks.append(
                        enrich_protocol_metadata(
                            {
                                "page_num": sheet_index + 1,
                                "content": content,
                                "type": "table",
                                "metadata": {"sheet_name": sheet.title},
                            },
                            False,
                        )
                    )
                _emit_progress(
                    progress_callback,
                    stage="quick_sampling_excel",
                    message=f"已采样工作表 {offset}/{len(sample_indexes)}",
                    progress=15.0 + (offset / max(len(sample_indexes), 1)) * 65.0,
                    processed_pages=offset,
                    total_pages=len(sample_indexes),
                )
        finally:
            workbook.close()
        total_units = len(sheet_names)
        sampled_indexes = [index + 1 for index in sample_indexes]
    else:
        import xlrd

        workbook = xlrd.open_workbook(file_path)
        total_units = workbook.nsheets
        sample_indexes = build_sample_indexes(total_units, sample_unit_limit)
        for offset, sheet_index in enumerate(sample_indexes, start=1):
            sheet = workbook.sheet_by_index(sheet_index)
            rows = []
            for row_index in range(min(sheet.nrows, EXCEL_SAMPLE_ROWS)):
                values = [
                    str(sheet.cell_value(row_index, col_index)).strip()
                    for col_index in range(min(sheet.ncols, TABLE_CELL_LIMIT))
                ]
                cleaned = [value for value in values if value]
                if cleaned:
                    rows.append(" | ".join(cleaned))
            content = _truncate_text("\n".join(rows))
            if content:
                blocks.append(
                    enrich_protocol_metadata(
                        {
                            "page_num": sheet_index + 1,
                            "content": content,
                            "type": "table",
                            "metadata": {"sheet_name": sheet.name},
                        },
                        False,
                    )
                )
            _emit_progress(
                progress_callback,
                stage="quick_sampling_excel",
                message=f"已采样工作表 {offset}/{len(sample_indexes)}",
                progress=15.0 + (offset / max(len(sample_indexes), 1)) * 65.0,
                processed_pages=offset,
                total_pages=len(sample_indexes),
            )
        sampled_indexes = [index + 1 for index in sample_indexes]

    if total_units <= 0:
        issues.append("未识别到有效工作表")
    if not blocks:
        issues.append("采样工作表未提取到可读内容")
    return _build_metrics(
        total_units=total_units,
        blocks=blocks,
        sampled_units=len(sampled_indexes),
        sampled_unit_indexes=sampled_indexes,
    ), issues


def _sample_text_content(file_path: str) -> Tuple[Dict[str, Any], List[str]]:
    issues: List[str] = []
    with open(file_path, "r", encoding="utf-8", errors="ignore") as handle:
        content = _truncate_text(handle.read(DEFAULT_TEXT_CHAR_LIMIT))
    blocks: List[Dict[str, Any]] = []
    if content:
        blocks.append(
            enrich_protocol_metadata(
                {
                    "page_num": 1,
                    "content": content,
                    "type": "text",
                    "metadata": {},
                },
                False,
            )
        )
    else:
        issues.append("采样文本为空或不可读")
    return _build_metrics(
        total_units=1 if content else 0,
        blocks=blocks,
        sampled_units=1 if content else 0,
        sampled_unit_indexes=[1] if content else [],
    ), issues


def analyze_file_content_quick(
    file_path: str,
    file_name: str,
    *,
    progress_callback: Optional[Callable[[dict], None]] = None,
    sample_unit_limit: int = DEFAULT_QUICK_SAMPLE_UNITS,
) -> Tuple[Dict[str, Any], List[str]]:
    ext = Path(file_name or file_path).suffix.lower()
    resolved_limit = _normalize_sample_limit(sample_unit_limit)
    if ext == ".pdf":
        return _sample_pdf_content(
            file_path,
            progress_callback=progress_callback,
            sample_unit_limit=resolved_limit,
        )
    if ext == ".docx":
        return _sample_docx_content(
            file_path,
            progress_callback=progress_callback,
            sample_unit_limit=resolved_limit,
        )
    if ext in {".xlsx", ".xls"}:
        return _sample_excel_content(
            file_path,
            ext,
            progress_callback=progress_callback,
            sample_unit_limit=resolved_limit,
        )
    _emit_progress(
        progress_callback,
        stage="quick_sampling_text",
        message="已完成文本采样校验",
        progress=80.0,
    )
    return _sample_text_content(file_path)
