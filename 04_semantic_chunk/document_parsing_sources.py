from __future__ import annotations

from pathlib import Path
from typing import List, Tuple

from database.models import Block
from document_parsing_common import (
    append_block,
    build_direct_index_project_id,
    derive_file_names_from_document_paths,
    format_table_rows,
    split_text_chunks,
    validate_document_paths,
)
from document_parsing_pdf import append_pdf_blocks


def append_docx_blocks(file_path: str, project_id: str, file_name: str, blocks: List[Block]) -> None:
    from docx import Document

    document = Document(file_path)
    current_text: List[str] = []
    logical_page = 0

    for element in document.element.body:
        if element.tag.endswith("tbl"):
            if current_text:
                text_content = "\n".join(current_text).strip()
                if text_content:
                    logical_page += 1
                    for chunk in split_text_chunks(text_content):
                        append_block(
                            blocks,
                            project_id=project_id,
                            file_name=file_name,
                            page_num=logical_page,
                            content=chunk,
                            block_type="text",
                            source_document_path=file_path,
                        )
                current_text = []

            table = next((tbl for tbl in document.tables if tbl._tbl == element), None)
            if table is None:
                continue
            table_rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            table_text = format_table_rows(table_rows)
            if not table_text:
                continue
            logical_page += 1
            append_block(
                blocks,
                project_id=project_id,
                file_name=file_name,
                page_num=logical_page,
                content=table_text,
                block_type="table",
                source_document_path=file_path,
                extra_metadata={
                    "row_count": len(table_rows),
                    "col_count": max((len(row) for row in table_rows), default=0),
                },
            )
            continue

        if element.tag.endswith("p"):
            paragraph = next((para for para in document.paragraphs if para._p == element), None)
            if paragraph is not None and paragraph.text.strip():
                current_text.append(paragraph.text.strip())

    if current_text:
        logical_page += 1
        text_content = "\n".join(current_text).strip()
        for chunk in split_text_chunks(text_content):
            append_block(
                blocks,
                project_id=project_id,
                file_name=file_name,
                page_num=logical_page,
                content=chunk,
                block_type="text",
                source_document_path=file_path,
            )


def append_excel_blocks(file_path: str, project_id: str, file_name: str, blocks: List[Block]) -> None:
    ext = Path(file_path).suffix.lower()
    sheets: List[Tuple[int, str, List[List[str]]]] = []

    if ext == ".xlsx":
        from openpyxl import load_workbook

        workbook = load_workbook(file_path, data_only=True, read_only=True)
        try:
            for sheet_index, sheet in enumerate(workbook.worksheets, start=1):
                rows: List[List[str]] = []
                for row in sheet.iter_rows(values_only=True):
                    cleaned = [str(cell).strip() if cell is not None else "" for cell in row]
                    if any(cleaned):
                        rows.append(cleaned)
                sheets.append((sheet_index, sheet.title, rows))
        finally:
            workbook.close()
    elif ext == ".xls":
        import xlrd

        workbook = xlrd.open_workbook(file_path)
        for sheet_index in range(workbook.nsheets):
            sheet = workbook.sheet_by_index(sheet_index)
            rows = []
            for row_idx in range(sheet.nrows):
                cleaned = [str(sheet.cell_value(row_idx, col_idx)).strip() for col_idx in range(sheet.ncols)]
                if any(cleaned):
                    rows.append(cleaned)
            sheets.append((sheet_index + 1, sheet.name, rows))
    else:
        raise ValueError(f"不支持的Excel文件类型: {ext}")

    for page_num, sheet_name, rows in sheets:
        table_text = format_table_rows(rows)
        if not table_text:
            continue
        append_block(
            blocks,
            project_id=project_id,
            file_name=file_name,
            page_num=page_num,
            content=table_text,
            block_type="table",
            source_document_path=file_path,
            extra_metadata={
                "sheet_name": sheet_name,
                "sheet_index": page_num,
                "row_count": len(rows),
                "col_count": max((len(row) for row in rows), default=0),
            },
        )


def append_text_blocks(file_path: str, project_id: str, file_name: str, blocks: List[Block]) -> None:
    text = Path(file_path).read_text(encoding="utf-8", errors="ignore")
    for index, chunk in enumerate(split_text_chunks(text), start=1):
        append_block(
            blocks,
            project_id=project_id,
            file_name=file_name,
            page_num=index,
            content=chunk,
            block_type="text",
            source_document_path=file_path,
        )


def load_blocks_from_document_paths(document_paths: List[str], project_id_hint: str = "") -> Tuple[str, List[Block]]:
    resolved_paths = validate_document_paths(document_paths)
    file_names = derive_file_names_from_document_paths(resolved_paths)
    project_id = build_direct_index_project_id(project_id_hint, file_names, resolved_paths)
    blocks: List[Block] = []

    for file_path in resolved_paths:
        file_name = Path(file_path).name.strip() or "document"
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            append_pdf_blocks(file_path, project_id, file_name, blocks)
        elif ext == ".docx":
            append_docx_blocks(file_path, project_id, file_name, blocks)
        elif ext in {".xlsx", ".xls"}:
            append_excel_blocks(file_path, project_id, file_name, blocks)
        else:
            append_text_blocks(file_path, project_id, file_name, blocks)

    if not blocks:
        raise ValueError("上传文件未解析出可用于建立索引的文本内容")
    return project_id, blocks
